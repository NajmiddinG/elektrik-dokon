from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponse, FileResponse
from main_app.forms import LoginForm
from django.shortcuts import redirect, render, get_object_or_404
from django.contrib import messages
from django.utils import timezone
from django.db.models import Min, Max, Q

from .models import User, Worker, WorkDay
from ishchi_app.models import Work, WorkAmount, WorkDayMoney, Money
from obyekt_app.models import Obyekt, WorkAmount, Given_money
from dokon_app.models import Product, ProductType, HistoryCame, HistoryObject, HistorySoldOut, ProductHistoryCame, ProductHistoryObject, ProductHistorySoldOut

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


def check_user(request):
    try:
        work = request.user.workers.exists()
        user = request.user
        if not user or not work:
            return False
        return True

    except Exception as e:
        print(2, e)
        return False


def check_user_type(request):
    user = request.user
    try:
        admin = user.workers.filter(name__iexact='admin').first()
        if admin:
            response = redirect('main_app:dashboard')
            response.set_cookie('user', str(user.id))
            response.set_cookie('worker', str(admin.id))
            return response
        dokon_worker = user.workers.filter(Q(name__iexact='dokon')).first()
        if dokon_worker:
            response = redirect('dokon_app:dashboard')
            response.set_cookie('user', str(user.id))
            response.set_cookie('worker', str(dokon_worker.id))
            return response
        obyekt_worker = user.workers.filter(Q(name__iexact='obyekt')).first()
        if obyekt_worker:
            response = redirect('obyekt_app:dashboard')
            response.set_cookie('user', str(user.id))
            response.set_cookie('worker', str(obyekt_worker.id))
            return response
        ishchi_worker = user.workers.filter(Q(name__iexact='ishchi')).first()
        if ishchi_worker:
            response = redirect('ishchi_app:dashboard')
            response.set_cookie('user', str(user.id))
            response.set_cookie('worker', str(ishchi_worker.id))
            return response
        logout(request)
        return redirect(user_login)
    except Exception as e:
        print(1, e)
        return HttpResponse('Tizimda xatolik')


def has_some_error(request):
    try:
        if request.user.workers.filter(name__iexact='admin').first(): return False
    except: pass
    return bool(not request.user.is_authenticated or not check_user(request))


def user_login(request):
    if request.user.is_authenticated:
        return check_user_type(request)
    
    context = {}
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            user = authenticate(request,
                                username = data['login'],
                                password = data['password'])
            if user is not None:
                if user.is_active:
                    login(request, user)
                    return check_user_type(request)
                else:
                    return HttpResponse('Kirmadingiz')
            else:
                return HttpResponse('Login yoki parolda xatolik bor')
    else:
        form = LoginForm()
        context = {
            'form': form,
        }
    return render(request, 'registration/login.html', context)


def logout_user(request):
    logout(request)
    return redirect('/login/')


def dashboard(request):
    if has_some_error(request): return redirect('/login/')
    users_with_workers_info = User.objects.filter(workers__isnull=False).values(
        'id',
        'username',
        'first_name',
        'last_name',
        'password',
        'tel_number',
        'address',
        'workers__name'
    )
    users_with_workers_info = []

    for user_info in User.objects.filter(workers__isnull=False).values(
        'id',
        'username',
        'first_name',
        'last_name',
        'password',
        'tel_number',
        'address',
        'workers__name',
    ):
        is_working = WorkDay.objects.filter(responsible_id=user_info['id'], end_date__isnull=True).exists()
        user_info['is_working'] = is_working
        users_with_workers_info.append(user_info)
    worker_type = request.user.workers.values_list('name', flat=True).first()
    context = {
        'active': 'main_1',
        'worker_type': worker_type,
        'users_with_workers_info': users_with_workers_info,
        'users_type': Worker.objects.all().values('name'),
    }
    response = render(request, 'main_app/user.html', context=context)
    return response


def create_user(request):
    if has_some_error(request): return redirect('/login/')
    if request.method=='POST':
        try:
            username = request.POST.get('username')
            first_name = request.POST.get('first_name')
            last_name = request.POST.get('last_name')
            tel_number = request.POST.get('tel_number')
            address = request.POST.get('address')
            password = request.POST.get('password')
            worker_type = request.POST.get('workers')

            # Create a new user instance
            user = User.objects.create_user(
                username=username,
                first_name=first_name,
                last_name=last_name,
                tel_number=tel_number,
                address=address,
            )

            user.set_password(password)

            # Find the existing worker
            existing_worker = Worker.objects.get(name=worker_type)
            existing_worker.user.add(user)
            user.save()
            existing_worker.save()
            messages.success(request, f'{first_name} muvaffaqiyatli yaratildi.')
            response = redirect('main_app:dashboard')
            return response
        except:
            messages.error(request, 'Xatolik yuz berdi.')
    return redirect('main_app:dashboard')


def edit_user(request, user_id):
    if has_some_error(request): return redirect('/login/')
    if request.method == 'POST':
        try:
            if bool(request.user.workers.filter(name__iexact='admin').first()):
                user = User.objects.get(id=user_id)
                user.username = str(request.POST.get('username'))
                user.first_name = str(request.POST.get('first_name')).capitalize()
                user.last_name = str(request.POST.get('last_name')).capitalize()
                user.tel_number = str(request.POST.get('tel_number')).capitalize()
                user.address = str(request.POST.get('address')).capitalize()
                if len(str(request.POST.get('password'))):
                    user.set_password(str(request.POST.get('password')))
                user.save()

                new_user_work = request.POST.get('workers')
                new_worker = get_object_or_404(Worker, name=new_user_work)

                old_user = user.workers.first()  # Assuming a WorkAmount can be associated with only one Obyekt
                if old_user:
                    old_user.user.remove(user)
                new_worker.user.add(user)

                messages.success(request, f"{user.first_name} ma'lumotlari muvaffaqiyatli o\'zgartirildi.")
                response = redirect('main_app:dashboard')
                return response
        except User.DoesNotExist:
            messages.error(request, 'Bunday id ga ega foydalanuvchi mavjud emas.')

    return redirect('main_app:dashboard')


def start_job(request):
    if has_some_error(request): return redirect('/login/')
    referer = request.META.get('HTTP_REFERER')
    if referer and not WorkDay.objects.filter(responsible=request.user, end_date__isnull=True).exists():
        WorkDay.objects.create(responsible=request.user)
        return redirect(referer)
    else:
        return redirect('/')


def end_job(request):
    if has_some_error(request): return redirect('/login/')
    referer = request.META.get('HTTP_REFERER')
    if referer and WorkDay.objects.filter(responsible=request.user, end_date__isnull=True).exists():
        WorkDay.objects.filter(responsible=request.user, end_date__isnull=True).update(end_date=timezone.now())
        return redirect(referer)
    else:
        return redirect('/')


def set_cookie_for_all_types_of_filter_view(request, name, value):
    if has_some_error(request): 
        return redirect('/login/')

    response = redirect(request.META.get('HTTP_REFERER', '/'))
    response.set_cookie(str(name), str(value))

    return response


def create_obyekt_worker_months(request):
    if has_some_error(request): return redirect('/login/')
    if request.method=='POST':
        try:
            cookies = request.COOKIES
            selected_obyekt1 = int(cookies.get('worker_admin_id', 0))
            selected_obyekt2 = int(cookies.get('worker_date_admin_id', 24290))

            name = str(request.POST.get('name'))
            money = int(request.POST.get('money'))
            given_money = Money.objects.create(
                responsible_id=selected_obyekt1,
                name=name,
                given_amount=money,
                month=selected_obyekt2
            )
            given_money.save()
            messages.success(request, f"{money} pull miqdori {User.objects.get(id=selected_obyekt1).first_name}  ga muaffaqiyatli qo'shildi!")
        except Exception as e:
            print(e)
            messages.error(request, 'Xatolik yuz berdi.')

    return redirect('main_app:ishchilar_holati')


def ishchilar_holati(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    cookies = request.COOKIES
    selected_obyekt1 = int(cookies.get('worker_admin_id', 0))
    if selected_obyekt1==0:
        for user in User.objects.all():
            if user.workers.filter(name__iexact='ishchi').first():
                selected_obyekt1 = user.id
                break
    selected_obyekt2 = int(cookies.get('worker_date_admin_id', 24289))
    try:
        first_date = WorkDayMoney.objects.filter(responsible_id=selected_obyekt1).aggregate(Min('date'))['date__min']
        last_date = WorkDayMoney.objects.filter(responsible_id=selected_obyekt1).aggregate(Max('date'))['date__max']
        first_one = 12*first_date.year+first_date.month
        last_one = 12*last_date.year+last_date.month
        months = [i for i in range(first_one, last_one+1)]
    except:
        months = []
    try:
        workdaymoneys_obyekt = WorkDayMoney.objects.filter(
            responsible_id=selected_obyekt1,
            date__year=(selected_obyekt2-1) // 12,
            date__month=(selected_obyekt2-1) % 12+1
        ).order_by('-date')
        month_given_amount = Money.objects.filter(responsible_id=selected_obyekt1,  month=selected_obyekt2)
        workdaymoneys2 = []
        for detail1 in workdaymoneys_obyekt:
            workdaymoneys = {}
            workdaymoneys['id']=detail1.id
            workdaymoneys['responsible']=detail1.responsible
            workdaymoneys['earn_amount']=detail1.earn_amount
            workdaymoneys['work_amount']=detail1.work_amount
            workdaymoneys['date']=detail1.date
            work_amount2 = detail1.work_amount.first().job
            obyekt_data = list(work_amount2.obyekt_set.values().first().values())
            workdaymoneys['obyekt_id'] = obyekt_data[0]
            workdaymoneys['obyekt_name'] = obyekt_data[2]
            workdaymoneys2.append(workdaymoneys)
        workdaymoneys = workdaymoneys2
    except Exception as e:
        print(e)
        workdaymoneys = WorkDayMoney.objects.filter(responsible_id=selected_obyekt1).order_by('-date')
        month_given_amount = Money.objects.filter(responsible_id=selected_obyekt1)
    
    work_money_earn = 0
    for workdaymoney_item in workdaymoneys:
        work_money_earn += workdaymoney_item['earn_amount']
    
    given_total = 0
    for item in month_given_amount:
        given_total+=item.given_amount
    
    obyekt_workers = User.objects.filter(workers__name='Obyekt')
    worker_type = request.user.workers.values_list('name', flat=True).first()
    context = {
        'active': 'main_2',
        'month_given_amount': month_given_amount,
        'obyekt_workers': obyekt_workers,
        'worker_type': worker_type,
        'work_money_earn': work_money_earn,
        'workdaymoneys': workdaymoneys,
        'given_total': given_total,
        'all_workers': Worker.objects.get(name='Ishchi').user.all(),
        'months': months,
        'real_money': given_total-work_money_earn
    }
    response = render(request, 'main_app/ishchilar_holati.html', context=context)
    try:
        response.set_cookie('worker_admin_id', str(selected_obyekt1))
        response.set_cookie('worker_date_admin_id', str(selected_obyekt2))
    except Exception as e:
        print(e)
    return response


def obyekt_material(request):
    if has_some_error(request): return redirect('/login/')

    cookies = request.COOKIES
    sold_history_id = int(cookies.get('sold_history_id', 0))
    obyekt_id_report = int(cookies.get('obyekt_id_report', 0))
    histoysoldouts = HistoryObject.objects.filter(history_object__id=obyekt_id_report).order_by('-date')
    obyekts = Obyekt.objects.all().order_by('-date')
    try:
        products = HistoryObject.objects.get(id=sold_history_id).history_products.all()
    except:
        products = []
    cur_date = timezone.now()
    year, month = cur_date.year, cur_date.month

    months = [i for i in range(2024*12+1, year*12+month+1)]
    context = {
        'active': 'main_5',
        'histoysoldouts': histoysoldouts,
        'months': months,
        'obyekts': obyekts,
        'products': products,
        'worker_type': request.user.workers.values_list('name', flat=True).first()
    }
    print(histoysoldouts)
    response = render(request, 'main_app/material.html', context=context)
    if sold_history_id == 0:
        response.set_cookie('sold_history_id', '0')
    return response


def change_materials(request):
    if has_some_error(request): return redirect('/login/')
    
    cookies = request.COOKIES
    sold_history_id = int(cookies.get('sold_history_id', 0))
    history_obj = HistoryObject.objects.get(id=sold_history_id)
    try:
        total_money = 0
        sold_out_products = []
        history_sold_outs = [request.user, 0, 0, 0]
        conflict_number = 0
        conflict_money = 0
        conflict_profit = 0
        for key, number in request.POST.items():
            if key.startswith('quantity;') and number!='0':
                pro_id = int(key.split(';')[1])
                pro_obj = ProductHistoryObject.objects.get(id=pro_id)
                old = pro_obj.number
                qaytib_keldi = int(number)
                if qaytib_keldi==0: continue
                new = old-qaytib_keldi
                pro_obj.number = new
                conflict_number += qaytib_keldi
                conflict_money += pro_obj.total_amount
                pro_obj.total_amount = (pro_obj.total_amount//old)*new
                conflict_money -= pro_obj.total_amount
                conflict_profit += pro_obj.profit
                pro_obj.profit = (pro_obj.profit//old)*new
                conflict_profit -= pro_obj.profit
                pro_obj.save()

        # Create a HistorySoldOut object
        history_obj.total_amount -= conflict_money
        history_obj.total_number_given -= conflict_number
        history_obj.profit -= conflict_profit
        obyekt = history_obj.history_object.id
        obyekt = Obyekt.objects.get(id=obyekt)
        obyekt.real_dept += conflict_money
        obyekt.save()
        history_obj.save()


        messages.success(request, f" Obyektdan material qaytarish muvaffaqiyatli amalga oshirildi!")

    except Exception as e:
        print(e, 2324234)
        messages.error(request, "Xatolik ro'y berdi!")
    return redirect('main_app:obyekt_material')


def obyekt_material_obyekt(request):
    if has_some_error(request): return redirect('/login/')

    cookies = request.COOKIES
    sold_history_id = int(cookies.get('sold_history_id', 0))
    obyekt_id_report = int(cookies.get('obyekt_id_report', 0))
    histoysoldouts = HistoryObject.objects.filter(history_object__id=obyekt_id_report).order_by('-date')
    if bool(request.user.workers.filter(name__iexact='admin').first()):
        obyekts = Obyekt.objects.all().order_by('-date')
    else:
        obyekts = Obyekt.objects.filter(responsible=request.user).order_by('-date')
    try:
        products = HistoryObject.objects.get(id=sold_history_id).history_products.all()
    except:
        products = []
    cur_date = timezone.now()
    year, month = cur_date.year, cur_date.month

    months = [i for i in range(2024*12+1, year*12+month+1)]
    context = {
        'active': 'main_6',
        'histoysoldouts': histoysoldouts,
        'months': months,
        'obyekts': obyekts,
        'products': products,
        'worker_type': request.user.workers.values_list('name', flat=True).first()
    }
    response = render(request, 'main_app/material2.html', context=context)
    if sold_history_id == 0:
        response.set_cookie('sold_history_id', '0')
    return response


def edit_obyekt_worker_months(request, done_work_id):
    if has_some_error(request): return redirect('/login/')
    if request.method=='POST':
        try:
            work_id = int(request.POST.get('work_id'))
            total_completed = int(request.POST.get('total_completed'))
            work_obj = Work.objects.get(id=work_id)
            diff = total_completed-work_obj.completed
            diff_price_worker = diff*work_obj.job.service_price
            diff_price_object = diff*work_obj.job.first_price
            # changing Work according to diff
            work_obj.completed = total_completed
            work_obj.save()
            # changing workdayamount according to diff
            workdaymoney_obj = work_obj.workdaymoney_set.values().first()
            workdaymoney_obj = WorkDayMoney.objects.get(id=workdaymoney_obj['id'])
            workdaymoney_obj.earn_amount += diff_price_worker
            workdaymoney_obj.save()
            # changing obyekt's work amount according to diff
            work_amount_obj = WorkAmount.objects.get(id=work_obj.job.id)
            work_amount_obj.total_completed+=diff
            work_amount_obj.save()
            # changing obyekt according to diff
            obyekt_obj = work_amount_obj.obyekt_set.values().first()
            obyekt_obj = Obyekt.objects.get(id=obyekt_obj['id'])
            obyekt_obj.real_dept-=diff_price_object
            obyekt_obj.save()

            messages.success(request, "Obyekt ni real qarzdorligi, Obyekt ishining tugatilganlar soni, Qilingan ish qiymati va Qilingan ishni soni muvaffaqiyatli o'zgardi!")
        except Exception as e:
            print(e)
            messages.error(request, 'Xatolik yuz berdi.')

    return redirect('main_app:done_work_detail')


def done_work_detail(request):
    if has_some_error(request): return redirect('/login/')

    cookies = request.COOKIES
    selected_obyekt = int(cookies.get('workdaymoney_id', 0))

    cookies = request.COOKIES
    selected_obyekt1 = int(cookies.get('worker_admin_id', 0))
    if selected_obyekt1==0:
        for user in User.objects.all():
            if user.workers.filter(name__iexact='ishchi').first():
                selected_obyekt1 = user.id
                break
    selected_obyekt2 = int(cookies.get('worker_date_admin_id', 24289))
    try:
        first_date = WorkDayMoney.objects.filter(responsible_id=selected_obyekt1).aggregate(Min('date'))['date__min']
        last_date = WorkDayMoney.objects.filter(responsible_id=selected_obyekt1).aggregate(Max('date'))['date__max']
        first_one = 12*first_date.year+first_date.month
        last_one = 12*last_date.year+last_date.month
        months = [i for i in range(first_one, last_one+1)]
    except:
        months = []
    try:
        if selected_obyekt == 0:
            work_amounts = []
            for entry in WorkDayMoney.objects.filter(responsible_id=selected_obyekt1, date__year=(selected_obyekt2-1) // 12,
                date__month=(selected_obyekt2-1) % 12+1).order_by('-date'):
                w = entry.work_amount.all()
                for i in range(len(w)):
                    w[i].date = entry.date
                work_amounts.extend(list(w))
        else: 
            a = WorkDayMoney.objects.get(id=selected_obyekt, responsible_id=selected_obyekt1)
            work_amounts = a.work_amount.all()
            for i in range(len(work_amounts)):
                work_amounts[i].date = a.date

    except:
        work_amounts = []

    try:
        workdaymoneys = WorkDayMoney.objects.filter(responsible_id=selected_obyekt1, 
            date__year=(selected_obyekt2-1) // 12,
            date__month=(selected_obyekt2-1) % 12+1).order_by('-date')
    except Exception as e:
        print(e)
        workdaymoneys = []
    
    
    worker_type = request.user.workers.values_list('name', flat=True).first()
    context = {
        'active': 'main_3',
        'workdaymoneys': workdaymoneys,
        'worker_type': worker_type,
        'work_amounts': work_amounts,
        'all_workers': Worker.objects.get(name='Ishchi').user.all(),
        'months': months,
    }

    response = render(request, 'main_app/ishchilar_ishlari.html', context=context)
    try:
        if not work_amounts and WorkDayMoney.objects.get(id=selected_obyekt):
            response.set_cookie('workdaymoney_id', str(0))
        response.set_cookie('worker_admin_id', str(selected_obyekt1))
        response.set_cookie('worker_date_admin_id', str(selected_obyekt2))
    except Exception as e:
        print(e)
    return response


def hisobotlar(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    cur_date = timezone.now()
    year, month = cur_date.year, cur_date.month

    months = [i for i in range(2024*12+1, year*12+month+1)]
    days = [i for i in range(1, 32)]
    obyekts = Obyekt.objects.all()
    worker_type = request.user.workers.values_list('name', flat=True).first()
    context = {
        'active': 'main_4',
        'worker_type': worker_type,
        'months': months,
        'days': days,
        'obyekts': obyekts,
        'all_workers': Worker.objects.get(name='Ishchi').user.all(),
    }
    response = render(request, 'main_app/hisobotlar.html', context=context)
    return response


def month_accurate_format(value):
    value = int(value)
    months_table = {
        1:'Yanvar',
        2:'Fevral',
        3:'Mart',
        4:'Aprel',
        5:'May',
        6:'Iyun',
        7:'Iyul',
        8:'Avgust',
        9:'Sentyabr',
        10:'Oktabr',
        11:'Noyabr',
        0:'Dekabr',

    }
    return f'{months_table[value%12]} - {(value-1)//12}'


def spacecomma(value):
    s_text = ''
    money = str(value)[::-1]
    for i in range(len(money)//3+1):
        s_text+=money[3*i:3*i+3]+' '
    return s_text.strip()[::-1] + " so'm"


def bool_to_word(value):
    if value: return "Ha"
    return "Yo'q"


def today_accurate_format():
    today = datetime.today()
    return today.strftime('%d-%m-%Y')


def day_format(date):
    return date.strftime('%d-%m-%Y')


def day_format2(day, month):
    try:
        date = datetime(day=day, month=(month-1)%12+1, year=(month-1)//12)
        return date.strftime('%d-%m-%Y')
    except:
        return 'Xato oy raqami'


def obyekt_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    obyekt_id = int(request.COOKIES.get('obyekt_id_report', 0))
    obyekt = Obyekt.objects.get(id=obyekt_id)
    day_id = int(request.COOKIES.get('seleted_day_id', 1))
    
    history_solt_out = HistoryObject.objects.filter(
                history_object=obyekt
            ).order_by('-date')
    products = {}
    for item in history_solt_out:
        for item2 in item.history_products.all():
            if item2.id in products:
                products[item2.id]['number'] += item2.number
                products[item2.id]['amount'] += item2.total_amount
                products[item2.id]['profit'] += item2.profit
            else:
                products[item2.id] = {}
                products[item2.id]['obyekt'] = item.history_object.name
                products[item2.id]['category'] = item2.type.type.first_type
                products[item2.id]['type'] = item2.type.type.name
                products[item2.id]['name'] = item2.type.name
                products[item2.id]['price'] = item2.type.price
                products[item2.id]['number'] = item2.number
                products[item2.id]['amount'] = item2.total_amount
                products[item2.id]['profit'] = item2.profit

    month_report = {'total_amount': 0, 'total_number_sold_out': 0, 'profit': 0}
    for item in products.values():
        month_report['total_amount'] += item['amount']
        month_report['total_number_sold_out'] += item['number']
        month_report['profit'] += item['profit']

    given_money = Given_money.objects.filter(obyekt=obyekt)

    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"{obyekt.name} obyekt {today_accurate_format()} hisoboti.", level=1)
    doc.add_heading("Obyekt hisoboti", level=2)
    table = doc.add_table(rows=1, cols=8)
    table.width = Inches(6.5)
    widths = [1, 2, 2, 2, 2, 2, 2, 2]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Nomi'
    hdr_cells[2].text = 'Masul'
    hdr_cells[3].text = 'Ish turi'
    hdr_cells[4].text = 'Kelishilgan summa'
    hdr_cells[5].text = 'Olingan summa'
    hdr_cells[6].text = 'Real qarzdorlik'
    hdr_cells[7].text = 'Bajarilgan ish'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in range(1):
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = obyekt.name
        row_cells[2].text = obyekt.responsible.first_name
        row_cells[3].text = obyekt.job_type
        row_cells[4].text = spacecomma(obyekt.deal_amount)
        row_cells[5].text = spacecomma(obyekt.given_amount)
        row_cells[6].text = spacecomma(obyekt.real_dept)
        row_cells[7].text = spacecomma(obyekt.given_amount-obyekt.real_dept)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    

    doc.add_heading("Berilgan summa hisoboti", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.width = Inches(6.5)
    widths = [2, 5, 3, 5, 5]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Masul'
    hdr_cells[2].text = 'Olingan summa'
    hdr_cells[3].text = 'Izoh'
    hdr_cells[4].text = 'Sana'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in given_money:
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item.responsible.first_name
        row_cells[2].text = spacecomma(item.amount)
        row_cells[3].text = item.comment
        row_cells[4].text = day_format(item.date)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.add_heading("Qilingan ishlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.width = Inches(6.5)
    widths = [1.5, 5, 3, 3, 3, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Ish turi'
    hdr_cells[2].text = 'Admin narx'
    hdr_cells[3].text = 'Xizmat narx'
    hdr_cells[4].text = 'Tugagan soni'
    hdr_cells[5].text = 'Umumiy soni'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in obyekt.work_amount.all():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item.job_type
        row_cells[2].text = spacecomma(item.first_price)
        row_cells[3].text = spacecomma(item.service_price)
        row_cells[4].text = str(item.total_completed)
        row_cells[5].text = str(item.total)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.add_heading("Materiallar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [1.2, 5, 3, 5]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy savdo'
    hdr_cells[2].text = 'Sotilgan mahsulotlar soni'
    hdr_cells[3].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in range(1):
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(month_report['total_amount'])
        row_cells[2].text = str(month_report['total_number_sold_out'])
        row_cells[3].text = spacecomma(month_report['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.add_heading("Obyektga biriktirilgan mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=9)
    table.width = Inches(6.5)
    widths = [2, 4, 4, 4, 3, 3, 3, 3, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Obyekt'
    hdr_cells[2].text = 'Categoriya'
    hdr_cells[3].text = 'Turi'
    hdr_cells[4].text = 'Nomi'
    hdr_cells[5].text = 'Narxi'
    hdr_cells[6].text = 'Soni'
    hdr_cells[7].text = 'Savdo'
    hdr_cells[8].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for id, item in products.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item['obyekt']
        row_cells[2].text = item['category']
        row_cells[3].text = item['type']
        row_cells[4].text = item['name']
        row_cells[5].text = spacecomma(item['price'])
        row_cells[6].text = str(item['number'])
        row_cells[7].text = spacecomma(item['amount'])
        row_cells[8].text = spacecomma(item['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Obyekt {today_accurate_format()} sana hisoboti.docx")


def history_sold_out_day_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    month = int(request.COOKIES.get('worker_date_admin_id', 24289))
    day_id = int(request.COOKIES.get('seleted_day_id', 1))
    
    history_solt_out = HistorySoldOut.objects.filter(
                date__year=(month-1) // 12,
                date__month=(month-1) % 12+1,
                date__day=day_id,
            ).order_by('-date')
    products = {}
    for item in history_solt_out:
        for item2 in item.history_products.all():
            if item2.id in products:
                products[item2.id]['number'] += item2.number
                products[item2.id]['amount'] += item2.total_amount
                products[item2.id]['profit'] += item2.profit
            else:
                products[item2.id] = {}
                products[item2.id]['category'] = item2.type.type.first_type
                products[item2.id]['type'] = item2.type.type.name
                products[item2.id]['name'] = item2.type.name
                products[item2.id]['price'] = item2.type.price
                products[item2.id]['number'] = item2.number
                products[item2.id]['amount'] = item2.total_amount
                products[item2.id]['profit'] = item2.profit

    month_report = {'total_amount': 0, 'total_number_sold_out': 0, 'profit': 0}
    for item in products.values():
        month_report['total_amount'] += item['amount']
        month_report['total_number_sold_out'] += item['number']
        month_report['profit'] += item['profit']


    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Do'kon {day_format2(day_id, month)} hisoboti.", level=1)
    doc.add_heading("Kunlik hisobot", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [1.2, 5, 3, 5]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy savdo'
    hdr_cells[2].text = 'Sotilgan mahsulotlar soni'
    hdr_cells[3].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in range(1):
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(month_report['total_amount'])
        row_cells[2].text = str(month_report['total_number_sold_out'])
        row_cells[3].text = spacecomma(month_report['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.add_heading("Kunlik mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=8)
    table.width = Inches(6.5)
    widths = [2, 4, 4, 4, 3, 3, 3, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Categoriya'
    hdr_cells[2].text = 'Turi'
    hdr_cells[3].text = 'Nomi'
    hdr_cells[4].text = 'Narxi'
    hdr_cells[5].text = 'Soni'
    hdr_cells[6].text = 'Savdo'
    hdr_cells[7].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for id, item in products.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item['category']
        row_cells[2].text = item['type']
        row_cells[3].text = item['name']
        row_cells[4].text = spacecomma(item['price'])
        row_cells[5].text = str(item['number'])
        row_cells[6].text = spacecomma(item['amount'])
        row_cells[7].text = spacecomma(item['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Savdo {day_format2(day_id, month)} hisobot.docx")


def history_sold_out_to_obyekt_day_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    month = int(request.COOKIES.get('worker_date_admin_id', 24289))
    day_id = int(request.COOKIES.get('seleted_day_id', 1))
    
    history_solt_out = HistoryObject.objects.filter(
                date__year=(month-1) // 12,
                date__month=(month-1) % 12+1,
                date__day=day_id,
            ).order_by('-date')
    products = {}
    for item in history_solt_out:
        for item2 in item.history_products.all():
            if item2.id in products:
                products[item2.id]['number'] += item2.number
                products[item2.id]['amount'] += item2.total_amount
                products[item2.id]['profit'] += item2.profit
            else:
                products[item2.id] = {}
                products[item2.id]['obyekt'] = item.history_object.name
                products[item2.id]['category'] = item2.type.type.first_type
                products[item2.id]['type'] = item2.type.type.name
                products[item2.id]['name'] = item2.type.name
                products[item2.id]['price'] = item2.type.price
                products[item2.id]['number'] = item2.number
                products[item2.id]['amount'] = item2.total_amount
                products[item2.id]['profit'] = item2.profit

    month_report = {'total_amount': 0, 'total_number_sold_out': 0, 'profit': 0}
    for item in products.values():
        month_report['total_amount'] += item['amount']
        month_report['total_number_sold_out'] += item['number']
        month_report['profit'] += item['profit']


    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Do'konni obyektga bergan mahsulotlar {day_format2(day_id, month)} hisoboti.", level=1)    
    doc.add_heading("Kunlik hisobot", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [1.2, 5, 3, 5]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy savdo'
    hdr_cells[2].text = 'Sotilgan mahsulotlar soni'
    hdr_cells[3].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in range(1):
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(month_report['total_amount'])
        row_cells[2].text = str(month_report['total_number_sold_out'])
        row_cells[3].text = spacecomma(month_report['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    
    doc.add_heading("Kunlik mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=9)
    table.width = Inches(6.5)
    widths = [2, 4, 4, 4, 3, 3, 3, 3, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Obyekt'
    hdr_cells[2].text = 'Categoriya'
    hdr_cells[3].text = 'Turi'
    hdr_cells[4].text = 'Nomi'
    hdr_cells[5].text = 'Narxi'
    hdr_cells[6].text = 'Soni'
    hdr_cells[7].text = 'Savdo'
    hdr_cells[8].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for id, item in products.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item['obyekt']
        row_cells[2].text = item['category']
        row_cells[3].text = item['type']
        row_cells[4].text = item['name']
        row_cells[5].text = spacecomma(item['price'])
        row_cells[6].text = str(item['number'])
        row_cells[7].text = spacecomma(item['amount'])
        row_cells[8].text = spacecomma(item['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1




    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Dokon-obyekt {day_format2(day_id, month)} sana hisoboti.docx")


def history_came_day_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    month = int(request.COOKIES.get('worker_date_admin_id', 24289))
    day_id = int(request.COOKIES.get('seleted_day_id', 1))
    
    history_solt_out = HistoryCame.objects.filter(
                date__year=(month-1) // 12,
                date__month=(month-1) % 12+1,
                date__day=day_id,
            ).order_by('-date')
    products = {}
    for item in history_solt_out:
        for item2 in item.history_products.all():
            if item2.id in products:
                products[item2.id]['number'] += item2.number
                products[item2.id]['amount'] += item2.total_amount
            else:
                products[item2.id] = {}
                products[item2.id]['category'] = item2.type.type.first_type
                products[item2.id]['type'] = item2.type.type.name
                products[item2.id]['name'] = item2.type.name
                products[item2.id]['price'] = item2.type.price
                products[item2.id]['number'] = item2.number
                products[item2.id]['amount'] = item2.total_amount

    month_report = {'total_amount': 0, 'total_number_sold_out': 0}
    for item in products.values():
        month_report['total_amount'] += item['amount']
        month_report['total_number_sold_out'] += item['number']


    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Do'konga kelgan mahsulotlar {day_format2(day_id, month)} hisoboti.", level=1)
    doc.add_heading("Kunlik hisobot", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.width = Inches(6.5)
    widths = [1, 7, 7]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy kelgan tovar'
    hdr_cells[2].text = 'Kelgan mahsulotlar soni'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in range(1):
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(month_report['total_amount'])
        row_cells[2].text = str(month_report['total_number_sold_out'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.add_heading("Kunlik mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.width = Inches(6.5)
    widths = [1.5, 4, 4, 4, 3, 3, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Categoriya'
    hdr_cells[2].text = 'Turi'
    hdr_cells[3].text = 'Nomi'
    hdr_cells[4].text = 'Narxi'
    hdr_cells[5].text = 'Soni'
    hdr_cells[6].text = 'Umumiy narx'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for id, item in products.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item['category']
        row_cells[2].text = item['type']
        row_cells[3].text = item['name']
        row_cells[4].text = spacecomma(item['price'])
        row_cells[5].text = str(item['number'])
        row_cells[6].text = spacecomma(item['amount'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Do'konga kelgan mahsulotlar {day_format2(day_id, month)} sana hisoboti.docx")


def history_sold_out_products_current_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    month = int(request.COOKIES.get('worker_date_admin_id', 24289))
    history_solt_out = HistorySoldOut.objects.filter(
                date__year=(month-1) // 12,
                date__month=(month-1) % 12+1
            ).order_by('-date')
    category_month = {}
    products = {}
    for item in history_solt_out:
        if item.date in category_month:
            category_month[day_format(item.date)]['total_amount'] += item.total_amount
            category_month[day_format(item.date)]['total_number_sold_out'] += item.total_number_sold_out
            category_month[day_format(item.date)]['profit'] += item.profit
        else:
            category_month[day_format(item.date)] = {}
            category_month[day_format(item.date)]['total_amount'] = item.total_amount
            category_month[day_format(item.date)]['total_number_sold_out'] = item.total_number_sold_out
            category_month[day_format(item.date)]['profit'] = item.profit
        for item2 in item.history_products.all():
            if item2.id in products:
                products[item2.id]['number'] += item2.number
                products[item2.id]['amount'] += item2.total_amount
                products[item2.id]['profit'] += item2.profit
            else:
                products[item2.id] = {}
                products[item2.id]['category'] = item2.type.type.first_type
                products[item2.id]['type'] = item2.type.type.name
                products[item2.id]['name'] = item2.type.name
                products[item2.id]['price'] = item2.type.price
                products[item2.id]['number'] = item2.number
                products[item2.id]['amount'] = item2.total_amount
                products[item2.id]['profit'] = item2.profit

    month_report = {'total_amount': 0, 'total_number_sold_out': 0, 'profit': 0}
    for item in category_month.values():
        month_report['total_amount'] += item['total_amount']
        month_report['total_number_sold_out'] += item['total_number_sold_out']
        month_report['profit'] += item['profit']


    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Do'kon hisoboti.", level=1)
    doc.add_heading("Oylik hisobot", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [1.2, 5, 3, 5]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy savdo'
    hdr_cells[2].text = 'Sotilgan mahsulotlar soni'
    hdr_cells[3].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in range(1):
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(month_report['total_amount'])
        row_cells[2].text = str(month_report['total_number_sold_out'])
        row_cells[3].text = spacecomma(month_report['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    
    doc.add_heading("Kunlik hisobot", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.width = Inches(6.5)
    widths = [1, 5, 3, 5, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy savdo'
    hdr_cells[2].text = 'Soni'
    hdr_cells[3].text = 'Foyda'
    hdr_cells[4].text = 'Sana'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for day, item in category_month.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(item['total_amount'])
        row_cells[2].text = str(item['total_number_sold_out'])
        row_cells[3].text = spacecomma(item['profit'])
        row_cells[4].text = day

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.add_heading("Oylik mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=8)
    table.width = Inches(6.5)
    widths = [2, 4, 4, 4, 3, 3, 3, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Categoriya'
    hdr_cells[2].text = 'Turi'
    hdr_cells[3].text = 'Nomi'
    hdr_cells[4].text = 'Narxi'
    hdr_cells[5].text = 'Soni'
    hdr_cells[6].text = 'Savdo'
    hdr_cells[7].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for id, item in products.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item['category']
        row_cells[2].text = item['type']
        row_cells[3].text = item['name']
        row_cells[4].text = spacecomma(item['price'])
        row_cells[5].text = str(item['number'])
        row_cells[6].text = spacecomma(item['amount'])
        row_cells[7].text = spacecomma(item['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1




    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Savdo {today_accurate_format()} sana hisoboti.docx")


def history_sold_out_to_obyekt_products_current_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    month = int(request.COOKIES.get('worker_date_admin_id', 24289))
    history_solt_out = HistoryObject.objects.filter(
                date__year=(month-1) // 12,
                date__month=(month-1) % 12+1
            ).order_by('-date')
    category_month = {}
    products = {}
    for item in history_solt_out:
        if item.date in category_month:
            category_month[day_format(item.date)]['total_amount'] += item.total_amount
            category_month[day_format(item.date)]['total_number_sold_out'] += item.total_number_given
            category_month[day_format(item.date)]['profit'] += item.profit
        else:
            category_month[day_format(item.date)] = {}
            category_month[day_format(item.date)]['total_amount'] = item.total_amount
            category_month[day_format(item.date)]['total_number_sold_out'] = item.total_number_given
            category_month[day_format(item.date)]['profit'] = item.profit
        for item2 in item.history_products.all():
            if item2.id in products:
                products[item2.id]['number'] += item2.number
                products[item2.id]['amount'] += item2.total_amount
                products[item2.id]['profit'] += item2.profit
            else:
                products[item2.id] = {}
                products[item2.id]['obyekt'] = item.history_object.name
                products[item2.id]['category'] = item2.type.type.first_type
                products[item2.id]['type'] = item2.type.type.name
                products[item2.id]['name'] = item2.type.name
                products[item2.id]['price'] = item2.type.price
                products[item2.id]['number'] = item2.number
                products[item2.id]['amount'] = item2.total_amount
                products[item2.id]['profit'] = item2.profit

    month_report = {'total_amount': 0, 'total_number_sold_out': 0, 'profit': 0}
    for item in category_month.values():
        month_report['total_amount'] += item['total_amount']
        month_report['total_number_sold_out'] += item['total_number_sold_out']
        month_report['profit'] += item['profit']


    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Do'konni obyektga bergan mahsulotlar hisoboti.", level=1)
    doc.add_heading("Oylik hisobot", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [1.2, 5, 3, 5]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy savdo'
    hdr_cells[2].text = 'Sotilgan mahsulotlar soni'
    hdr_cells[3].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in range(1):
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(month_report['total_amount'])
        row_cells[2].text = str(month_report['total_number_sold_out'])
        row_cells[3].text = spacecomma(month_report['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    
    doc.add_heading("Kunlik hisobot", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.width = Inches(6.5)
    widths = [1, 5, 3, 5, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy savdo'
    hdr_cells[2].text = 'Soni'
    hdr_cells[3].text = 'Foyda'
    hdr_cells[4].text = 'Sana'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for day, item in category_month.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(item['total_amount'])
        row_cells[2].text = str(item['total_number_sold_out'])
        row_cells[3].text = spacecomma(item['profit'])
        row_cells[4].text = day

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.add_heading("Oylik mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=9)
    table.width = Inches(6.5)
    widths = [2, 4, 4, 4, 3, 3, 3, 3, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Obyekt'
    hdr_cells[2].text = 'Categoriya'
    hdr_cells[3].text = 'Turi'
    hdr_cells[4].text = 'Nomi'
    hdr_cells[5].text = 'Narxi'
    hdr_cells[6].text = 'Soni'
    hdr_cells[7].text = 'Savdo'
    hdr_cells[8].text = 'Foyda'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for id, item in products.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item['obyekt']
        row_cells[2].text = item['category']
        row_cells[3].text = item['type']
        row_cells[4].text = item['name']
        row_cells[5].text = spacecomma(item['price'])
        row_cells[6].text = str(item['number'])
        row_cells[7].text = spacecomma(item['amount'])
        row_cells[8].text = spacecomma(item['profit'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1




    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Dokon-obyekt {today_accurate_format()} sana hisoboti.docx")


def history_came_products_current_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    month = int(request.COOKIES.get('worker_date_admin_id', 24289))
    history_solt_out = HistoryCame.objects.filter(
                date__year=(month-1) // 12,
                date__month=(month-1) % 12+1
            ).order_by('-date')
    category_month = {}
    products = {}
    for item in history_solt_out:
        if item.date in category_month:
            category_month[day_format(item.date)]['total_amount'] += item.total_amount
            category_month[day_format(item.date)]['total_number_sold_out'] += item.total_number_sold_out
        else:
            category_month[day_format(item.date)] = {}
            category_month[day_format(item.date)]['total_amount'] = item.total_amount
            category_month[day_format(item.date)]['total_number_sold_out'] = item.total_number_sold_out
        for item2 in item.history_products.all():
            if item2.id in products:
                products[item2.id]['number'] += item2.number
                products[item2.id]['amount'] += item2.total_amount
            else:
                products[item2.id] = {}
                products[item2.id]['category'] = item2.type.type.first_type
                products[item2.id]['type'] = item2.type.type.name
                products[item2.id]['name'] = item2.type.name
                products[item2.id]['price'] = item2.type.price
                products[item2.id]['number'] = item2.number
                products[item2.id]['amount'] = item2.total_amount

    month_report = {'total_amount': 0, 'total_number_sold_out': 0}
    for item in category_month.values():
        month_report['total_amount'] += item['total_amount']
        month_report['total_number_sold_out'] += item['total_number_sold_out']


    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Do'konga kelgan mahsulotlar hisoboti.", level=1)
    doc.add_heading("Oylik hisobot", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.width = Inches(6.5)
    widths = [1, 5, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy kelgan tovar'
    hdr_cells[2].text = 'Kelgan mahsulotlar soni'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for item in range(1):
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(month_report['total_amount'])
        row_cells[2].text = str(month_report['total_number_sold_out'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    
    doc.add_heading("Kunlik hisobot", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [1, 5, 3, 5]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy kelgan tovar'
    hdr_cells[2].text = 'Soni'
    hdr_cells[3].text = 'Sana'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for day, item in category_month.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = spacecomma(item['total_amount'])
        row_cells[2].text = str(item['total_number_sold_out'])
        row_cells[3].text = day

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.add_heading("Oylik mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.width = Inches(6.5)
    widths = [1.5, 4, 4, 4, 3, 3, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Categoriya'
    hdr_cells[2].text = 'Turi'
    hdr_cells[3].text = 'Nomi'
    hdr_cells[4].text = 'Narxi'
    hdr_cells[5].text = 'Soni'
    hdr_cells[6].text = 'Umumiy narx'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    index_t = 0
    for id, item in products.items():
        index_t+=1
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = item['category']
        row_cells[2].text = item['type']
        row_cells[3].text = item['name']
        row_cells[4].text = spacecomma(item['price'])
        row_cells[5].text = str(item['number'])
        row_cells[6].text = spacecomma(item['amount'])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Do'konga kelgan mahsulotlar {today_accurate_format()} sana hisoboti.docx")


def elektrik_products_current_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Do'konning {today_accurate_format()} hisoboti.", level=1)
    doc.add_heading("Elektr mahsulotlar", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.width = Inches(6.5)
    widths = [2, 4, 4, 2, 2, 2, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Turi'
    hdr_cells[2].text = 'Nomi'
    hdr_cells[3].text = 'Kelish narx'
    hdr_cells[4].text = 'Sotilish narx'
    hdr_cells[5].text = 'Qoldiq'
    hdr_cells[6].text = 'Umumiy narx'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    total_kelish = 0
    total_sotish = 0
    index_t = 0
    all_products = Product.objects.filter(type__first_type='elektr').order_by('type__name')
    print(all_products.count())
    for product in all_products.iterator():
        print(index_t)
        index_t+=1
        sotish_narx = (product.price*(100+product.profit_percentage))//100
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = product.type.name
        row_cells[2].text = product.name
        row_cells[3].text = spacecomma(product.price)
        row_cells[4].text = spacecomma(sotish_narx)
        row_cells[5].text = str(product.remain)
        row_cells[6].text = spacecomma(sotish_narx * product.remain)
        total_kelish += product.price*product.remain
        total_sotish += sotish_narx*product.remain

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.add_heading("Umumiy Elektr mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=3)
    widths = [1,3,3]
    table.width = Inches(6.5)
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy tan narx'
    hdr_cells[2].text = 'Umumiy mahsulot miqdori'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = str(1)
    row_cells[1].text = spacecomma(total_kelish)
    row_cells[2].text = spacecomma(total_sotish)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Elektr mahsulotlari {today_accurate_format()} sana hisoboti.docx")


def santexnika_products_current_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Do'konning {today_accurate_format()} hisoboti.", level=1)
    doc.add_heading("Santexnika mahsulotlar", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.width = Inches(6.5)
    widths = [2, 4, 4, 2, 2, 2, 3]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Turi'
    hdr_cells[2].text = 'Nomi'
    hdr_cells[3].text = 'Kelish narx'
    hdr_cells[4].text = 'Sotilish narx'
    hdr_cells[5].text = 'Qoldiq'
    hdr_cells[6].text = 'Umumiy narx'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    total_kelish = 0
    total_sotish = 0
    index_t = 0
    all_products = Product.objects.filter(type__first_type='santexnika').order_by('type__name')
    print(all_products.count())
    for product in all_products.iterator():
        print(index_t)
        index_t+=1
        sotish_narx = (product.price*(100+product.profit_percentage))//100
        row_cells = table.add_row().cells
        row_cells[0].text = str(index_t)
        row_cells[1].text = product.type.name
        row_cells[2].text = product.name
        row_cells[3].text = spacecomma(product.price)
        row_cells[4].text = spacecomma(sotish_narx)
        row_cells[5].text = str(product.remain)
        row_cells[6].text = spacecomma(sotish_narx * product.remain)
        total_kelish += product.price*product.remain
        total_sotish += sotish_narx*product.remain

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.add_heading("Umumiy Santexnika mahsulotlar hisoboti", level=2)
    table = doc.add_table(rows=1, cols=3)
    widths = [1,3,3]
    table.width = Inches(6.5)
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Umumiy tan narx'
    hdr_cells[2].text = 'Umumiy mahsulot miqdori'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = str(1)
    row_cells[1].text = spacecomma(total_kelish)
    row_cells[2].text = spacecomma(total_sotish)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f"Santexnika mahsulotlari {today_accurate_format()} sana hisoboti.docx")


def create_monthly_workers_report(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    month = int(request.COOKIES.get('worker_date_admin_id', 24289))
    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Ishchilarning {month_accurate_format(month)} hisoboti.", level=1)
    doc.add_heading("Ishchining umumiy hisobot", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.width = Inches(6.5)
    widths = [2,7,5,5,5]
    table.width = Inches(6.5)
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Ishchi'
    hdr_cells[2].text = 'Berilgan summa'
    hdr_cells[3].text = 'Qilingan ish miqdori'
    hdr_cells[4].text = 'Qoldiq'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    total_1 = 0
    total_2 = 0
    index_t = 0
    for user in Worker.objects.get(name='Ishchi').user.all():
        index_t += 1
        moneys = Money.objects.filter(responsible=user, month=month)
        total_given = sum(money.given_amount for money in moneys)
        total_done_work_amount = 0

        try:
            workdaymoneys_obyekt = WorkDayMoney.objects.filter(
                responsible=user,
                date__year=(month-1) // 12,
                date__month=(month-1) % 12+1
            ).order_by('-date')
            total_done_work_amount += sum(money.earn_amount for money in workdaymoneys_obyekt)
        except Exception as e:
            workdaymoneys = []

        total_1 += total_given
        total_2 += total_done_work_amount

        for index in range(1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index_t)
            row_cells[1].text = user.first_name
            row_cells[2].text = spacecomma(total_given)
            row_cells[3].text = spacecomma(total_done_work_amount)
            row_cells[4].text = spacecomma(total_given - total_done_work_amount)

    for ind in range(1):
            index_t += 1
            row_cells = table.add_row().cells
            row_cells[0].text = str(index_t)
            row_cells[1].text = "Umumiy"
            row_cells[2].text = spacecomma(total_1)
            row_cells[3].text = spacecomma(total_2)
            row_cells[4].text = spacecomma(total_1 - total_2)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1  # 0 for top, 1 for center, 2 for bottom
    
    
    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f'{user.first_name}_{month_accurate_format(month)}.docx')


def generate_worker_pdf(request):
    if has_some_error(request) or not bool(request.user.workers.filter(name__iexact='admin').first()): return redirect('/login/')

    buffer = BytesIO()
    doc = Document()

    # Your existing code
    month = int(request.COOKIES.get('worker_date_admin_id', 24289))
    user_id = int(request.COOKIES.get('worker_admin_id', 0))
    user = User.objects.get(id=user_id)
    moneys = Money.objects.filter(responsible=user, month=month)
    total_given = sum(money.given_amount for money in moneys)
    total_done_work_amount = 0

    try:
        workdaymoneys_obyekt = WorkDayMoney.objects.filter(
            responsible=user,
            date__year=(month-1) // 12,
            date__month=(month-1) % 12+1
        ).order_by('-date')
        total_done_work_amount += sum(money.earn_amount for money in workdaymoneys_obyekt)
        workdaymoneys2 = []
        for detail1 in workdaymoneys_obyekt:
            workdaymoneys = {}
            workdaymoneys['id']=detail1.id
            workdaymoneys['responsible']=detail1.responsible
            workdaymoneys['earn_amount']=detail1.earn_amount
            workdaymoneys['work_amount']=detail1.work_amount
            workdaymoneys['date']=detail1.date
            work_amount2 = detail1.work_amount.first().job
            obyekt_data = list(work_amount2.obyekt_set.values().first().values())
            workdaymoneys['obyekt_id'] = obyekt_data[0]
            workdaymoneys['obyekt_name'] = obyekt_data[2]
            workdaymoneys2.append(workdaymoneys)
        workdaymoneys = workdaymoneys2
    except Exception as e:
        workdaymoneys = []

    doc.add_heading(f"{user.first_name} ning {month_accurate_format(month)} hisoboti.", level=1)
    doc.add_heading("Ishchining umumiy hisobot", level=2)

    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [2,5,5,5]
    table.width = Inches(6.5)
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Berilgan summa'
    hdr_cells[2].text = 'Qilingan ish miqdori'
    hdr_cells[3].text = 'Qoldiq'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for index in range(1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index+1)
        row_cells[1].text = spacecomma(total_given)
        row_cells[2].text = spacecomma(total_done_work_amount)
        row_cells[3].text = spacecomma(total_given - total_done_work_amount)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    doc.add_heading("Berilgan oylik hisoboti", level=2)

    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [2,5,5,5]
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Sabab'
    hdr_cells[2].text = 'Miqdor'
    hdr_cells[3].text = 'Berilgan sana'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for index, money in enumerate(moneys):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index+1)
        row_cells[1].text = money.name
        row_cells[2].text = str(spacecomma(money.given_amount))
        row_cells[3].text = money.date.strftime("%d-%m-%Y")

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1  # 0 for top, 1 for center, 2 for bottom
    
    doc.add_heading("Kunlik hisobotlar", level=2)

    table = doc.add_table(rows=1, cols=4)
    table.width = Inches(6.5)
    widths = [2,6,5,5]
    table.width = Inches(6.5)
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Obyekt nomi'
    hdr_cells[2].text = 'Miqdor'
    hdr_cells[3].text = 'Sana'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for index, workdaymoney in enumerate(workdaymoneys):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index+1)
        row_cells[1].text = workdaymoney['obyekt_name']
        row_cells[2].text = str(spacecomma(workdaymoney['earn_amount']))
        row_cells[3].text = workdaymoney['date'].strftime("%d-%m-%Y")

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.add_heading("Qilingan ishlar hisoboti", level=2)

    table = doc.add_table(rows=1, cols=5)
    table.width = Inches(6.5)
    widths = [2,7,5,2,5]
    table.width = Inches(6.5)
    total_widths = sum(widths)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width / total_widths * 6.5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Ish turi'
    hdr_cells[2].text = 'Xizmat haqqi'
    hdr_cells[3].text = 'Soni'
    hdr_cells[4].text = 'Sana'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    i = 0
    for index, workdaymoney in enumerate(workdaymoneys):
        for work in workdaymoney['work_amount'].all():
            i+=1
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = work.job.job_type
            row_cells[2].text = str(spacecomma(work.job.service_price))
            row_cells[3].text = str(work.completed)
            row_cells[4].text = workdaymoney['date'].strftime("%d-%m-%Y")

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
    
    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename=f'{user.first_name}_{month_accurate_format(month)}.docx')
