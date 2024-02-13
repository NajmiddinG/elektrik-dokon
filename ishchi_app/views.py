from django.contrib.auth import authenticate, login, logout
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.http import HttpResponse, JsonResponse
from main_app.views import has_some_error
from django.contrib import messages
from django.shortcuts import get_object_or_404
from django.db.models import Min, Max
from django.conf import settings
import os
from dokon_app.models import (
    ProductType,
    Product,
    ProductHistorySoldOut,
    HistorySoldOut,
    ProductHistoryCame,
    HistoryCame,
    ProductHistoryObject,
    HistoryObject,
    ObjectPayment,
)
from main_app.models import Worker, User, WorkDay
from django.contrib.auth import authenticate, login, logout
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.http import HttpResponse, JsonResponse
from main_app.views import has_some_error
from django.contrib import messages
from django.shortcuts import get_object_or_404

from obyekt_app.models import Obyekt, WorkAmount, WorkAmountJobType, ObyektJobType, Allow, Instructsiya, Obyekt_doc
from main_app.models import Worker, User, WorkDay
from ishchi_app.models import Work, WorkDayMoney, Money
from main_app.calculate import calculate_worker_to_obyekt

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from io import BytesIO
from django.http import FileResponse
import io

from docx import Document
from docx.shared import Pt


def dashboard(request):
    if has_some_error(request): return redirect('/login/')

    has_allow_entry = Allow.objects.filter(responsible=request.user).exists()
    instruktsiya_doc = Instructsiya.objects.all()
    is_working = WorkDay.objects.filter(responsible=request.user, end_date__isnull=True).exists()
    obyekts = Obyekt.objects.all().order_by('-date')
    obyekt_workers = User.objects.filter(workers__name='Obyekt')
    worker_type = request.user.workers.values_list('name', flat=True).first()
    obyektjobtypes = ObyektJobType.objects.all().order_by('name')
    workeramountjobtypes = WorkAmountJobType.objects.all().order_by('name')
    context = {
        'active': 'ishchi_1',
        'obyekts': obyekts,
        'obyekt_workers': obyekt_workers,
        'worker_type': worker_type,
        'obyektjobtypes': obyektjobtypes,
        'workeramountjobtypes': workeramountjobtypes,
        'position': 'end' if is_working else 'start',
        'allowed': has_allow_entry,
        'instruktsiya': instruktsiya_doc,

    }
    return render(request, 'ishchi/obyekt.html', context=context)

def obyekt_ishi(request):
    if has_some_error(request): return redirect('/login/')

    has_allow_entry = Allow.objects.filter(responsible=request.user).exists()
    instruktsiya_doc = Instructsiya.objects.all()
    cookies = request.COOKIES
    selected_obyekt = int(cookies.get('obyekt_id', 0))
    try:
        work_amounts = Obyekt.objects.get(pk=selected_obyekt).work_amount.all()
        obyekt_doc = Obyekt_doc.objects.filter(role='ishchi', obyekt=Obyekt.objects.get(pk=selected_obyekt))
    except:
        work_amounts = []
        obyekt_doc = []
    obyekts = Obyekt.objects.all().order_by('-date')
    is_working = WorkDay.objects.filter(responsible=request.user, end_date__isnull=True).exists()
    obyekt_workers = User.objects.filter(workers__name='Obyekt')
    worker_type = request.user.workers.values_list('name', flat=True).first()
    obyektjobtypes = ObyektJobType.objects.all().order_by('name')
    workeramountjobtypes = WorkAmountJobType.objects.all().order_by('name')
    context = {
        'active': 'ishchi_2',
        'obyekts': obyekts,
        'obyekt_workers': obyekt_workers,
        'worker_type': worker_type,
        'work_amounts': work_amounts,
        'obyektjobtypes': obyektjobtypes,
        'workeramountjobtypes': workeramountjobtypes,
        'allowed': has_allow_entry,
        'instruktsiya': instruktsiya_doc,
        'obyekt_doc': obyekt_doc,
        'position': 'end' if is_working else 'start',
    }
    response = render(request, 'ishchi/obyekt_ishi.html', context=context)
    if selected_obyekt==0:
        try:
            latest_obyekt = Obyekt.objects.latest('date').id
            response.set_cookie('obyekt_id', str(latest_obyekt))
        except:
            response.set_cookie('obyekt_id', '0')
    return response

def done_work_post(request):
    if has_some_error(request): return redirect('/login/')

    if request.method == 'POST':
        try:
            # 'quantity;1': ['0'], 'quantity;2': ['0'],
            print(request.POST)
            done_works = []
            history_came = [request.user, 0] # responsible, earn_amount
            for key, number in request.POST.items():
                if key.startswith('quantity;') and number!='0':
                    number = int(number)
                    product_id = int(key.split(';')[1])
                    product_details = {
                        'job_id': product_id,
                        'completed': number,
                    }
                    work_amount = WorkAmount.objects.get(id=product_id)
                    history_came[1]+=number*work_amount.service_price
                    done_works.append(product_details)
            if history_came[1]==0:
                messages.error(request, "Xatolik ro'y berdi!")
            else:
                history_came = WorkDayMoney.objects.create(
                    responsible=history_came[0],
                    earn_amount=history_came[1],
                )
                for product_details in done_works:
                    product_history = Work.objects.create(
                        job_id=product_details['job_id'],
                        completed=product_details['completed'],
                    )
                    
                    history_came.work_amount.add(product_history)
                calculate_worker_to_obyekt(history_came.id)
                messages.success(request, "Bajarilgan ishlaringiz muvaffaqqiyatli qo'shildi")
        except Exception as e:
            print(e)
            messages.error(request, "Xatolik ro'y berdi!")

    return redirect("ishchi_app:dashboard")
    

def done_work_list(request):
    if has_some_error(request): return redirect('/login/')
    try:
        first_date = WorkDayMoney.objects.filter(responsible=request.user).aggregate(Min('date'))['date__min']
        last_date = WorkDayMoney.objects.filter(responsible=request.user).aggregate(Max('date'))['date__max']
        first_one = 12*first_date.year+first_date.month
        last_one = 12*last_date.year+last_date.month
        months = [i for i in range(first_one, last_one+1)]
    except:
        months = []
    is_working = WorkDay.objects.filter(responsible=request.user, end_date__isnull=True).exists()
    cookies = request.COOKIES
    selected_obyekt = int(cookies.get('worker_months', 24289))
    if selected_obyekt in [0, 24287]:
        selected_obyekt = 24289
        
    try:
        month_given_amount = Money.objects.filter(responsible_id=request.user,  month=selected_obyekt)
        workdaymoneys_obyekt = WorkDayMoney.objects.filter(
            responsible=request.user,
            date__year=selected_obyekt // 12,
            date__month=selected_obyekt % 12
        ).order_by('-date')
        workdaymoneys2 = []
        for detail1 in workdaymoneys_obyekt:
            workdaymoneys = {}
            workdaymoneys['id']=detail1.id
            workdaymoneys['responsible']=detail1.responsible
            workdaymoneys['earn_amount']=detail1.earn_amount
            workdaymoneys['work_amount']=detail1.work_amount
            workdaymoneys['date']=detail1.date
            work_amount2 = detail1.work_amount.first().job
            obyekt_data = list(work_amount2.obyekt_set.values().first().values())
            workdaymoneys['obyekt_id'] = obyekt_data[0]
            workdaymoneys['obyekt_name'] = obyekt_data[2]
            workdaymoneys2.append(workdaymoneys)
            # print(workdaymoneys)
        workdaymoneys = workdaymoneys2
    except Exception as e:
        print(e)
        workdaymoneys = WorkDayMoney.objects.filter(responsible=request.user).order_by('-date')
        month_given_amount = Money.objects.filter(responsible_id=request.user)
    
    given_total = 0
    for item in month_given_amount:
        given_total+=item.given_amount

    work_money_earn = 0
    for workdaymoney_item in workdaymoneys:
        work_money_earn += workdaymoney_item['earn_amount']
    worker_type = request.user.workers.values_list('name', flat=True).first()
    context = {
        'active': 'ishchi_3',
        'workdaymoneys': workdaymoneys,
        'worker_type': worker_type,
        'position': 'end' if is_working else 'start',
        'months': months,
        'work_money_earn': work_money_earn,
        'month_given_amount': month_given_amount,
        'given_total': given_total,
        'real_money': given_total-work_money_earn,

    }
    response = render(request, 'ishchi/done_work_list.html', context=context)
    response.set_cookie('worker_months', selected_obyekt)
    return response

def done_work_detail(request):
    if has_some_error(request): return redirect('/login/')

    cookies = request.COOKIES
    selected_obyekt = int(cookies.get('workdaymoney_id', 0))
    try:
        work_amounts = WorkDayMoney.objects.get(id=selected_obyekt, responsible=request.user).work_amount.all()
    except:
        work_amounts = []
    workdaymoneys = WorkDayMoney.objects.filter(responsible=request.user).order_by('-date')
    worker_type = request.user.workers.values_list('name', flat=True).first()
    context = {
        'active': 'ishchi_4',
        'workdaymoneys': workdaymoneys,
        'worker_type': worker_type,
        'work_amounts': work_amounts,
    }
    response = render(request, 'ishchi/done_work_detail.html', context=context)
    if selected_obyekt==0:
        try:
            latest_obyekt = WorkDayMoney.objects.filter(responsible=request.user).latest('date').id
            response.set_cookie('workdaymoney_id', str(latest_obyekt))
        except:
            response.set_cookie('workdaymoney_id', '0')
    return response


def allow_add(request):
    if has_some_error(request):
        return redirect('/login/')
    
    if request.method == 'POST':
        try:
            Allow.objects.create(responsible=request.user)
            messages.success(request, 'Juda soz.')
        except Exception as e:
            messages.error(request, f'Xatolik yuz berdi: {str(e)}')

    referer = request.META.get('HTTP_REFERER')
    if referer:
        return redirect(referer)
    else:
        return redirect('/')

def month_accurate_format(value):
    value = int(value)
    months_table = {
        1:'Yanvar',
        2:'Fevral',
        3:'Mart',
        4:'Aprel',
        5:'May',
        6:'Iyun',
        7:'Iyul',
        8:'Avgust',
        9:'Sentyabr',
        10:'Oktabr',
        11:'Noyabr',
        0:'Dekabr',

    }
    return f'{months_table[value%12]} - {value//12}'

def spacecomma(value):
    s_text = ''
    money = str(value)[::-1]
    for i in range(len(money)//3+1):
        s_text+=money[3*i:3*i+3]+' '
    return s_text.strip()[::-1] + " so'm"

def bool_to_word(value):
    if value: return "Ha"
    return "Yo'q"
def generate_worker_pdf(request):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    textop = c.beginText()
    textop.setTextOrigin(inch, 10*inch)  # Adjust x and y coordinates as needed
    textop.setFont("Helvetica", 14)

    month = int(request.COOKIES.get('worker_months', 24289))

    moneys = Money.objects.filter(month=month)

    textop.setFont("Helvetica", 25)
    textop.textLine(f"{request.user.first_name} ning {month_accurate_format(month)} hisoboti.")
    textop.textLine()
    textop.setFont("Helvetica", 20)
    textop.textLine("Berilgan oylik hisoboti")
    textop.setFont("Helvetica", 14)
    lines = []
    for index, money in enumerate(moneys):
        lines.append("    Nomer: "+ str(index+1))
        lines.append("    Javobgar: "+ money.responsible.first_name)
        lines.append("    Sabab: "+ money.name)
        lines.append("    Miqdor: "+ str(spacecomma(money.given_amount)))
        lines.append("    Berilgan sana: "+ money.date.strftime("%d-%m-%Y"))
        lines.append("")
    
    try:
        workdaymoneys_obyekt = WorkDayMoney.objects.filter(
            responsible=request.user,
            date__year=month // 12,
            date__month=month % 12
        ).order_by('-date')
        workdaymoneys2 = []
        for detail1 in workdaymoneys_obyekt:
            workdaymoneys = {}
            workdaymoneys['id']=detail1.id
            workdaymoneys['responsible']=detail1.responsible
            workdaymoneys['earn_amount']=detail1.earn_amount
            workdaymoneys['work_amount']=detail1.work_amount
            workdaymoneys['date']=detail1.date
            work_amount2 = detail1.work_amount.first().job
            obyekt_data = list(work_amount2.obyekt_set.values().first().values())
            workdaymoneys['obyekt_id'] = obyekt_data[0]
            workdaymoneys['obyekt_name'] = obyekt_data[2]
            workdaymoneys2.append(workdaymoneys)
        workdaymoneys = workdaymoneys2
    except Exception as e:
        workdaymoneys = []

    lines.append(("Qilingan ishlar hisoboti", 20))

    for index, workdaymoney in enumerate(workdaymoneys):
        lines.append("    Nomer: "+ str(index+1))
        lines.append("    Javobgar: "+ workdaymoney['responsible'].first_name)
        lines.append("    Miqdor: "+ str(spacecomma(workdaymoney['earn_amount'])))
        lines.append("    Sana: "+ workdaymoney['date'].strftime("%d-%m-%Y"))
        lines.append("")
        lines.append(("    Qilingan ishlar", 18))
        for index, work_amount_item in enumerate(workdaymoney['work_amount'].all()):
            lines.append("         Nomer: "+ str(index+1))
            lines.append(f"        Ish turi: {work_amount_item.job.job_type}")
            lines.append(f"        Bo'glangan: {bool_to_word(int(work_amount_item.job.visible_obyekt))}")
            lines.append(f"        Narxi: {str(spacecomma(work_amount_item.job.service_price))}")
            lines.append(f"        Bajargan soni: {str(work_amount_item.completed)}")
            lines.append(f"        Ish haqqi: {str(spacecomma(work_amount_item.completed*work_amount_item.job.service_price))}")
            lines.append("")
    for line in lines:
        if type(line)==tuple:
            textop.setFont("Helvetica", line[1])
            textop.textLine(line[0])
            textop.setFont("Helvetica", 14)
        else:
            textop.textLine(line)

        if textop.getY() <  inch:
            c.drawText(textop)
            c.showPage()
            textop = c.beginText()
            textop.setTextOrigin(inch, 10*inch)

    c.drawText(textop)
    c.save()
    buffer.seek(0)  

    return FileResponse(buffer, as_attachment=True, filename=f'{request.user.first_name}:{month_accurate_format(month)}.pdf')




    # doc = Document()

    # # Define font sizes
    # font_sizes = {
    #     "default": Pt(14),
    #     "header": Pt(25),
    #     "subheader": Pt(20),
    #     "section_header": Pt(18),
    #     "section_subheader": Pt(16)
    # }

    # month = int(request.COOKIES.get('worker_months', 24289))
    # doc.add_heading(f"{request.user.first_name} ning {month_accurate_format(month)} hisoboti.", level=1)
    # doc.add_paragraph()
    # doc.add_paragraph("Berilgan oylik hisoboti")

    # moneys = Money.objects.filter(month=month)

    # for index, money in enumerate(moneys):
    #     p = doc.add_paragraph()
    #     p.add_run(f"Nomer: {index}")
    #     p.add_run(f"Javobgar: {money.responsible.first_name}")
    #     p.add_run(f"Sabab: {money.name}")
    #     p.add_run(f"Miqdor: {spacecomma(money.given_amount)}")
    #     p.add_run(f"Berilgan sana: {money.date.strftime('%d-%m-%Y')}")
    #     doc.add_paragraph()

    # doc.add_heading("Qilingan ishlar hisoboti", level=2)
    # doc.add_paragraph()

    # try:
    #     workdaymoneys_obyekt = WorkDayMoney.objects.filter(
    #         responsible=request.user,
    #         date__year=month // 12,
    #         date__month=month % 12
    #     ).order_by('-date')
    #     workdaymoneys2 = []
    #     for detail1 in workdaymoneys_obyekt:
    #         workdaymoneys = {}
    #         workdaymoneys['id']=detail1.id
    #         workdaymoneys['responsible']=detail1.responsible
    #         workdaymoneys['earn_amount']=detail1.earn_amount
    #         workdaymoneys['work_amount']=detail1.work_amount
    #         workdaymoneys['date']=detail1.date
    #         work_amount2 = detail1.work_amount.first().job
    #         obyekt_data = list(work_amount2.obyekt_set.values().first().values())
    #         workdaymoneys['obyekt_id'] = obyekt_data[0]
    #         workdaymoneys['obyekt_name'] = obyekt_data[2]
    #         workdaymoneys2.append(workdaymoneys)
    #     workdaymoneys = workdaymoneys2
    # except Exception as e:
    #     workdaymoneys = []
    # for index, workdaymoney in enumerate(workdaymoneys):
    #     p = doc.add_paragraph()
    #     p.add_run(f"Nomer: {index}\n")
    #     p.add_run(f"Javobgar: {workdaymoney['responsible'].first_name}\n")
    #     p.add_run(f"Miqdor: {spacecomma(workdaymoney['earn_amount'])}\n")
    #     p.add_run(f"Sana: {workdaymoney['date'].strftime('%d-%m-%Y')}\n")
    #     doc.add_paragraph()

    #     # Add "Qilingan ishlar" section
    #     doc.add_heading("Qilingan ishlar", level=3)
    #     doc.add_paragraph()

    #     # Add content from WorkAmount objects
    #     for work_amount_item in workdaymoney['work_amount'].all():
    #         p = doc.add_paragraph()
    #         p.add_run(f"Nomer: {index}\n")
    #         p.add_run(f"Ish turi: {work_amount_item.job.job_type}\n")
    #         p.add_run(f"Bo'glangan: {bool_to_word(int(work_amount_item.job.visible_obyekt))}\n")
    #         p.add_run(f"Narxi: {spacecomma(work_amount_item.job.service_price)}\n")
    #         p.add_run(f"Bajargan soni: {work_amount_item.completed}\n")
    #         p.add_run(f"Ish haqqi: {spacecomma(work_amount_item.completed * work_amount_item.job.service_price)}\n")
    #         doc.add_paragraph()

    # # Save the document
    # file_path = os.path.join(settings.MEDIA_ROOT, 'ishchi_doc', f"{request.user.first_name}.docx")

    # # Save the document to the specified file path
    # doc.save(file_path)

    # # Open the saved document as a binary file
    # with open(file_path, "rb") as f:
    #     # Create a FileResponse with the document content
    #     response = FileResponse(f, as_attachment=True, filename=f"{request.user.first_name}.docx")

    # # Return the FileResponse
    # return response
